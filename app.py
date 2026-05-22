import streamlit as st
import pandas as pd
from datetime import datetime, date, timedelta
import pytz
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import os
import re
import requests
 
# --- KONSTAN & MAPPING DATA ---
TEMPLATE_PKDS = [
    'PKD GOMBAK', 'PKD HULU LANGAT', 'PKD HULU SELANGOR', 'PKD KLANG',
    'PKD KUALA LANGAT', 'PKD KUALA SELANGOR', 'PKD PETALING',
    'PKD SABAK BERNAM', 'PKD SEPANG'
]
 
AVG_HARIAN_FIGURES = {
    "Denggi": 426, "COVID-19": 54, "HFMD": 52, "Tuberculosis": 28,
    "Keracunan Makanan": 22, "Measles": 12, "Viral Hepatitis": 9,
    "Avian Influenza": 8, "HIV/AIDS": 7, "Leptospirosis": 6,
    "Dysentry": 5, "Syphilis": 5, "Typhoid/Paratyphoid": 5,
    "Gonorrhoea": 2, "Pertussis": 2, "Malaria": 1, "Mers-Cov": 1
}
 
SHEET_ID = "1bjyNcntm-I6nRaIVkVdJqJRAzn5r2tYFfjUAN0emv9w"
GID = "0"
GSHEET_URL = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=csv&gid={GID}"
 
# --- URL GOOGLE SHEET BKK (RAW LINELISTING & JADUAL) ---
BKK_SPREADSHEET_ID = "1Fp6IORRfdWSJCTC8vqSSoQz6RpCpNXHzO6jj0tHEf2c"
URL_BKK_LINELISTING = f"https://docs.google.com/spreadsheets/d/{BKK_SPREADSHEET_ID}/export?format=csv&gid=1352807145"
URL_BKK_JADUAL = f"https://docs.google.com/spreadsheets/d/{BKK_SPREADSHEET_ID}/export?format=csv&gid=1342717767"
 
CHART_IMAGE_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vTDprYai1uaP1L-JP6kuHRZX18AmDHX0ROEzRE37DaCHMo0cNWUvRa8R-65RZAK7XFWI6pb_-X-jF24/pubchart?oid=1681812411&format=image"
 
# --- HELPERS ---
def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = parse_xml(r'<w:tblHeader {}/>'.format(nsdecls('w')))
    trPr.append(tblHeader)
 
def get_msia_time():
    msia_tz = pytz.timezone('Asia/Kuala_Lumpur')
    return datetime.now(msia_tz)
 
def format_penyakit_name(name):
    name_str = str(name).strip().upper()
    if any(x in name_str for x in ["HIV", "AIDS", "HFMD", "COVID-19"]):
        return name_str
    if "FOOD POISONING" in name_str:
        return "Keracunan Makanan"
    if name_str in ["DENGUE/DHF", "DENGUE"]:
        return "Denggi"
    return name_str.title()
 
def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hex_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)
 
def clean_val(val):
    if pd.isna(val) or str(val).strip() == "" or str(val).strip() == "-" or str(val).lower() == "nan":
        return "-"
    cleaned = re.sub(r'\s*\(.*?\)', '', str(val)).strip()
    return cleaned if cleaned != "" else "-"
 
def get_epi_week(target_date):
    start_date = date(2026, 1, 4)
    if target_date < start_date: return "N/A"
    days_diff = (target_date - start_date).days
    return f"{(days_diff // 7) + 1}/{target_date.year}"
 
def get_malay_date(target_date):
    days_ms = {"Monday": "Isnin", "Tuesday": "Selasa", "Wednesday": "Rabu", "Thursday": "Khamis", "Friday": "Jumaat", "Saturday": "Sabtu", "Sunday": "Ahad"}
    months_ms = {1: "Januari", 2: "Februari", 3: "Mac", 4: "April", 5: "Mei", 6: "Jun", 7: "Julai", 8: "Ogos", 9: "September", 10: "Oktober", 11: "November", 12: "Disember"}
    day_name = days_ms.get(target_date.strftime("%A"), "")
    month_name = months_ms.get(target_date.month, "")
    return f"{target_date.day:02d} {month_name} {target_date.year} ({day_name})"
 
def apply_font(run, size, bold=True):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
 
def add_table_title(doc, label, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    run_label = p.add_run(f"{label} : ")
    apply_font(run_label, 11, bold=True)
    run_title = p.add_run(title)
    apply_font(run_title, 11, bold=False)
    p.paragraph_format.space_after = Pt(6)
 
def add_pkd_note(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
     
    run_main = p.add_run("*Nota :")
    apply_font(run_main, 7, bold=True)
 
    senarai_daerah = [
        "GBK = Gombak",
        "HL = Hulu Langat",
        "HS = Hulu Selangor",
        "KLG = Klang",
        "KL = Kuala Langat",
        "KS = Kuala Selangor",
        "PTG = Petaling",
        "SB = Sabak Bernam",
        "SPG = Sepang"
    ]
     
    for daerah in senarai_daerah:
        p_item = doc.add_paragraph()
        p_item.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_item.paragraph_format.space_before = Pt(0)
        p_item.paragraph_format.space_after = Pt(2)
        p_item.paragraph_format.left_indent = Inches(0.4) 
         
        run_item = p_item.add_run(daerah)
        apply_font(run_item, 7, bold=False)
         
    p_item.paragraph_format.space_after = Pt(8)
 
# --- DOCX GENERATOR ---
def generate_docx(matrix_df, col_sums, wabak_df, vector_df, bkk_table_df, is_bkk_empty, bkk_details, df_yesterday_list):
    doc = Document()
    now_msia = get_msia_time()
    today = now_msia.date()
    yesterday = today - timedelta(days=1)
 
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(2.54)
    content_width = section.page_width - section.left_margin - section.right_margin
 
    logo_path = "logo.png.jpg"
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.8))
 
    titles = [
        "LAPORAN HARIAN KEJADIAN BENCANA, WABAK, KECEMASAN, KRISIS (BWKK)",
        "PUSAT KESIAPSIAGAAN DAN TINDAKCEPAT KRISIS (CPRC)",
        "JABATAN KESIHATAN NEGERI SELANGOR"
    ]
    for text in titles:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        apply_font(run, 10.5, bold=True)
        para.paragraph_format.space_after = Pt(0)
 
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
 
    info_table = doc.add_table(rows=1, cols=2)
    info_table.width = content_width 
    for i in range(2):
        cell = info_table.cell(0, i)
        set_cell_background(cell, "C6E0B4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 0:
            txt = f"\nTarikh : {get_malay_date(today)}\n(Sehingga jam 10.00 pagi)"
        else:
            txt = f"\nMinggu Epidemiologi : {get_epi_week(today)}"
        run = p.add_run(txt)
        apply_font(run, 11, bold=True)
 
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
 
    p1_head = doc.add_paragraph()
    apply_font(p1_head.add_run("1.0 Ringkasan Laporan Input Enotifikasi"), 11, bold=True)
     
    total_notifications = int(col_sums['Grand Total'])
    h11 = doc.add_paragraph()
    h11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
    h11_text = f"Jadual di bawah menunjukkan jumlah input e-Notifikasi di negeri Selangor. Sejumlah {total_notifications} input notifikasi telah diterima pada {get_malay_date(yesterday)} dengan pecahan mengikut penyakit seperti dalam Jadual 1.1."
    apply_font(h11.add_run(h11_text), 11, bold=False)
 
    add_table_title(doc, "Jadual 1.1", "Senarai Input e-Notifikasi")
    t1 = doc.add_table(rows=len(matrix_df) + 2, cols=len(TEMPLATE_PKDS) + 3)
    t1.style = 'Table Grid'
    t1.width = content_width 
     
    pkd_map = {'PKD GOMBAK': 'GBK', 'PKD HULU LANGAT': 'HL', 'PKD HULU SELANGOR': 'HS', 'PKD KLANG': 'KLG', 'PKD KUALA LANGAT': 'KL', 'PKD KUALA SELANGOR': 'KS', 'PKD PETALING': 'PTG', 'PKD SABAK BERNAM': 'SB', 'PKD SEPANG': 'SPG'}
     
    h_cells = t1.rows[0].cells
    for i in range(len(h_cells)):
        h_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        h_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
     
    apply_font(h_cells[0].paragraphs[0].add_run("Penyakit"), 8, bold=True)
    set_cell_background(h_cells[0], "BFDFFF")
     
    for i, pkd in enumerate(TEMPLATE_PKDS):
        cell = h_cells[i+1]
        apply_font(cell.paragraphs[0].add_run(pkd_map.get(pkd, pkd)), 8, bold=True)
        set_cell_background(cell, "BFDFFF")
     
    apply_font(h_cells[len(TEMPLATE_PKDS)+1].paragraphs[0].add_run("Jumlah"), 8, bold=True)
    set_cell_background(h_cells[len(TEMPLATE_PKDS)+1], "FFFF00")
    apply_font(h_cells[len(TEMPLATE_PKDS)+2].paragraphs[0].add_run("Average Harian"), 8, bold=True)
    set_cell_background(h_cells[len(TEMPLATE_PKDS)+2], "FFC000")
 
    for r_idx, (penyakit, row_data) in enumerate(matrix_df.iterrows()):
        row = t1.rows[r_idx + 1].cells
        nama_formatted = format_penyakit_name(penyakit)
        apply_font(row[0].paragraphs[0].add_run(nama_formatted), 8, bold=True)
        set_cell_background(row[0], "D9E9FF")
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
 
        for c_idx, pkd in enumerate(TEMPLATE_PKDS):
            cell = row[c_idx+1]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_font(p.add_run(str(int(row_data[pkd]))), 8, bold=True)
         
        gt_cell = row[len(TEMPLATE_PKDS)+1]
        gt_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        gt_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_font(gt_cell.paragraphs[0].add_run(str(int(row_data['Grand Total']))), 8, bold=True)
        set_cell_background(gt_cell, "FFFFB3")
         
        avg_cell = row[len(TEMPLATE_PKDS)+2]
        avg_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        avg_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_font(avg_cell.paragraphs[0].add_run(str(int(row_data.get('Average Harian', 0)))), 8, bold=True)
        set_cell_background(avg_cell, "FFC000")
 
    f_cells = t1.rows[-1].cells
    apply_font(f_cells[0].paragraphs[0].add_run("Jumlah"), 8, bold=True)
    set_cell_background(f_cells[0], "FFFF00")
    f_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
 
    for i, pkd in enumerate(TEMPLATE_PKDS):
        cell = f_cells[i+1]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_font(cell.paragraphs[0].add_run(str(int(col_sums[pkd]))), 8, bold=True)
        set_cell_background(cell, "FFFF00")
 
    f_cells[len(TEMPLATE_PKDS)+1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    f_cells[len(TEMPLATE_PKDS)+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_font(f_cells[len(TEMPLATE_PKDS)+1].paragraphs[0].add_run(str(int(col_sums['Grand Total']))), 8, bold=True)
    set_cell_background(f_cells[len(TEMPLATE_PKDS)+1], "FFFF00")
    set_cell_background(f_cells[len(TEMPLATE_PKDS)+2], "FFC000")
 
    add_pkd_note(doc)
 
    doc.add_page_break()
    p2_head = doc.add_paragraph()
    apply_font(p2_head.add_run("2.0 Ringkasan Laporan Notifikasi Wabak"), 11, bold=True)
     
    harian_total = int(wabak_df['HARIAN'].sum())
    h21 = doc.add_paragraph()
    h21.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
    h21_text = f"Jadual di bawah menunjukkan jumlah wabak harian, aktif dan kumulatif di negeri Selangor. Sejumlah {harian_total} input notifikasi wabak diterima pada {get_malay_date(yesterday)}."
    apply_font(h21.add_run(h21_text), 11, bold=False)
 
    add_table_title(doc, "Jadual 2.1", "Senarai Notifikasi Wabak")
    t2 = doc.add_table(rows=len(wabak_df) + 2, cols=4)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.autofit = False  
     
    col_widths_t2 = [content_width * 0.40, content_width * 0.20, content_width * 0.20, content_width * 0.20]
 
    for i, h in enumerate(["Penyakit", "Harian", "Aktif", "Kumulatif"]):
        cell = t2.cell(0, i)
        cell.width = col_widths_t2[i]
        apply_font(cell.paragraphs[0].add_run(h), 8, bold=True)
        set_cell_background(cell, "BFDFFF")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    for i, (penyakit, row_data) in enumerate(wabak_df.iterrows()):
        cells = t2.rows[i+1].cells
        for idx_w in range(4):
            cells[idx_w].width = col_widths_t2[idx_w]
             
        apply_font(cells[0].paragraphs[0].add_run(str(penyakit)), 8, bold=True)
        set_cell_background(cells[0], "D9E9FF")
        for idx, col_key in enumerate(['HARIAN', 'AKTIF', 'KUMULATIF'], start=1):
            run = cells[idx].paragraphs[0].add_run(str(int(row_data[col_key])))
            apply_font(run, 8, bold=True)
            cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    f2_cells = t2.rows[-1].cells
    footer_vals = ["Jumlah", str(int(wabak_df['HARIAN'].sum())), str(int(wabak_df['AKTIF'].sum())), str(int(wabak_df['KUMULATIF'].sum()))]
    for i, txt in enumerate(footer_vals):
        f2_cells[i].width = col_widths_t2[i]
        apply_font(f2_cells[i].paragraphs[0].add_run(txt), 8, bold=True)
        set_cell_background(f2_cells[i], "FFFF00")
        f2_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    doc.add_paragraph()
    tarikh_semalam_str = get_malay_date(yesterday)
    add_table_title(doc, "Jadual 2.2", f"Senarai Wabak Yang Dilaporkan pada {tarikh_semalam_str}")
     
    t21 = doc.add_table(rows=1, cols=5)
    t21.style = 'Table Grid'
    t21.width = content_width 
    t21.allow_autofit = False
    set_repeat_table_header(t21.rows[0])
 
    widths_21 = [content_width * 0.05, content_width * 0.2, content_width * 0.2, content_width * 0.4, content_width * 0.15]
    h21_headers = ["Bil", "Wabak", "Daerah", "Tempat Berlaku", "Kadar Serangan"]
    for i, txt in enumerate(h21_headers):
        cell = t21.cell(0, i)
        cell.width = widths_21[i]
        set_cell_background(cell, "BFDFFF")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_font(p.add_run(txt), 10, bold=True)
 
    if not df_yesterday_list:
        row = t21.add_row().cells
        row[0].merge(row[4])
        row[0].text = "Tiada wabak dilaporkan pada tarikh ini."
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        for idx, item in enumerate(df_yesterday_list, start=1):
            row = t21.add_row().cells
            for i in range(5): 
                row[i].width = widths_21[i]
                row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
             
            row[0].text = str(idx)
            p_wabak = row[1].paragraphs[0]
            p_wabak.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_name = p_wabak.add_run(str(item[0]))
            apply_font(run_name, 8, bold=False)
            p_wabak.add_run("\n")
            kategori_display = "(Household)" if str(item[3]).strip() == "Rumah Persendirian" else "(Institusi)"
            run_cat = p_wabak.add_run(kategori_display)
            apply_font(run_cat, 8, bold=False)
 
            row[2].text = str(item[1]).title() 
            row[3].text = str(item[2]) 
 
            n_kes = float(item[4]) if pd.notna(item[4]) else 0
            n_dedah = float(item[5]) if pd.notna(item[5]) else 0
             
            if n_dedah > 0:
                calc_pct = (n_kes / n_dedah) * 100
                pct_str = f"{int(calc_pct)}%" if calc_pct % 1 == 0 else f"{calc_pct:.2f}%"
            else:
                pct_str = "0%"
 
            p_ar = row[4].paragraphs[0]
            p_ar.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_ar_main = p_ar.add_run(f"{int(n_kes)}/{int(n_dedah)}")
            apply_font(run_ar_main, 8, bold=False)
            p_ar.add_run("\n")
            run_ar_pct = p_ar.add_run(f"({pct_str})")
            apply_font(run_ar_pct, 8, bold=False)
 
            for c in range(5):
                p = row[c].paragraphs[0]
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c == 3 else WD_ALIGN_PARAGRAPH.CENTER
                if p.runs: apply_font(p.runs[0], 8, bold=False)
 
    p3_head = doc.add_paragraph()
    p3_head.paragraph_format.page_break_before = True 
    apply_font(p3_head.add_run("3.0 Ringkasan Laporan Wabak Vektor"), 11, bold=True)
     
    try: 
        val_sum = float(vector_df.iloc[-1, 1]) + float(vector_df.iloc[-1, 3]) + float(vector_df.iloc[-1, 5])
        xx_v = int(val_sum)
    except: 
        xx_v = 0
     
    xx_v_display = "Tiada" if xx_v == 0 else str(xx_v)
    h31 = doc.add_paragraph()
    h31.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
    h31_text = f"Jadual di bawah menunjukkan jumlah wabak vektor harian dan kumulatif di negeri Selangor. Sebanyak {xx_v_display} notifikasi wabak vektor telah diterima pada {get_malay_date(yesterday)} dengan pecahan mengikut penyakit seperti dalam Jadual 3.1. Rajah 3.1 pula menunjukkan tren kes mingguan denggi yang didaftarkan dari tahun 2025 hingga kini."
    apply_font(h31.add_run(h31_text), 11, bold=False)
 
    add_table_title(doc, "Jadual 3.1", "Senarai Notifikasi Wabak Vektor")
    t3 = doc.add_table(rows=len(vector_df) + 2, cols=7)
    t3.style = 'Table Grid'
    t3.width = content_width 
    t3.allow_autofit = False  
 
    col_widths_v = [content_width * 0.25, content_width * 0.125, content_width * 0.125, content_width * 0.125, content_width * 0.125, content_width * 0.125, content_width * 0.125]
     
    h3_r1 = t3.rows[0].cells
    h3_r1[0].merge(t3.rows[1].cells[0]).text = "Daerah"
    h3_r1[1].merge(h3_r1[2]).text = "Denggi"
    h3_r1[3].merge(h3_r1[4]).text = "Malaria"
    h3_r1[5].merge(h3_r1[6]).text = "Chikungunya"
     
    for i in [0, 1, 3, 5]:
        cell = h3_r1[i]
        set_cell_background(cell, "BFDFFF")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if i == 0: cell.width = col_widths_v[0]
        else: cell.width = col_widths_v[i] + col_widths_v[i+1]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs: apply_font(p.runs[0], 10, bold=True)
        else: apply_font(p.add_run(cell.text), 10, bold=True)
 
    h3_r2 = t3.rows[1].cells
    for i in range(1, 7):
        h3_r2[i].text = "Harian" if i % 2 != 0 else "Kumulatif"
        h3_r2[i].width = col_widths_v[i]
        set_cell_background(h3_r2[i], "BFDFFF")
        h3_r2[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = h3_r2[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_font(p.runs[0], 9, bold=True)
 
    for i in range(len(vector_df)):
        row_cells = t3.rows[i+2].cells
        is_jumlah_row = str(vector_df.iloc[i, 0]).strip().upper() == "JUMLAH"
         
        for j in range(7):
            val = vector_df.iloc[i, j]
            row_cells[j].width = col_widths_v[j]
            if pd.isna(val) or str(val).lower() == "nan": display_val = "-"
            elif j == 0: 
                display_val = "Jumlah" if is_jumlah_row else str(val).title()
            else:
                try: display_val = f"{int(float(val)):,}"
                except: display_val = str(val)
             
            p = row_cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(display_val)
             
            if is_jumlah_row: 
                set_cell_background(row_cells[j], "FFFF00") 
            elif j == 0: 
                set_cell_background(row_cells[j], "FCE4D6") 
                 
            apply_font(run, 9, bold=True)
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
 
    try:
        response = requests.get(CHART_IMAGE_URL)
        if response.status_code == 200:
            doc.add_paragraph()
            p_rajah_head = doc.add_paragraph()
            p_rajah_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_rajah_label = p_rajah_head.add_run("Rajah 3.1 : ")
            apply_font(run_rajah_label, 11, bold=True)
            run_rajah_title = p_rajah_head.add_run("Tren Kes Mingguan Denggi Didaftar Bagi Tahun 2025 - 2026 Negeri Selangor")
            apply_font(run_rajah_title, 11, bold=False)
 
            img_stream = io.BytesIO(response.content)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(img_stream, width=Inches(6.2))
        else:
            st.warning("Gagal memuat turun imej graf.")
    except Exception as e:
        st.error(f"Ralat semasa memproses Rajah 3.1: {e}")
 
    # --- 4.0 Ringkasan Laporan Kejadian Insiden Bencana, Kecemasan dan Krisis (BKK) ---
    doc.add_page_break()
    p4_head = doc.add_paragraph()
    apply_font(p4_head.add_run("4.0 Ringkasan Laporan Kejadian Bencana, Kecemasan dan Krisis (BKK)"), 11, bold=True)
 
    h41 = doc.add_paragraph()
    h41.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
 
    penutup_text = " Jadual di bawah menunjukkan jumlah kejadian insiden bencana, kecemasan dan krisis (BKK) yang telah dilaporkan di negeri Selangor."
 
    if is_bkk_empty:
        h41_text = f"Tiada insiden dilaporkan pada {get_malay_date(yesterday)}.{penutup_text}"
    else:
        num_word = {1: "satu (1)", 2: "dua (2)", 3: "tiga (3)", 4: "empat (4)", 5: "lima (5)"}
        count = len(bkk_details)
        count_str = num_word.get(count, f"{count} ({count})")
         
        h41_text = f"Terdapat {count_str} kejadian dilaporkan pada {get_malay_date(yesterday)}."
         
        ordinal_words = {1: " Insiden pertama ialah ", 2: " Insiden kedua ialah ", 3: " Insiden ketiga ialah ", 4: " Insiden keempat ialah ", 5: " Insiden kelima ialah "}
         
        narrative_parts = []
        for idx, item in enumerate(bkk_details, start=1):
            prefix = ordinal_words.get(idx, f" Insiden ke-{idx} ialah ") if count > 1 else ""
             
            kej_str = str(item['kejadian']).lower()
            alamat_str = str(item['alamat']).strip()
            daerah_str = str(item['daerah']).title()
             
            def clean_to_int(v):
                if pd.isna(v) or str(v).strip() in ["", "-", "nan", "0"]:
                    return 0
                try:
                    return int(float(str(v).strip()))
                except:
                    return 0
 
            kes_val = clean_to_int(item.get('bil_kes', 0))
            kem_val = clean_to_int(item.get('bil_kematian', 0))
             
            if kes_val > 0:
                if kem_val > 0:
                    status_str = f" Sejumlah {kes_val} orang mangsa terlibat dalam kejadian {kej_str} tersebut dengan {kem_val} kematian dilaporkan."
                else:
                    status_str = f" Sejumlah {kes_val} orang mangsa terlibat dalam kejadian {kej_str} tersebut."
            else:
                if kem_val > 0:
                    status_str = f" Sejumlah {kem_val} kematian dilaporkan dalam kejadian {kej_str} tersebut."
                else:
                    status_str = ""
             
            if count == 1:
                full_event_sentence = f" Kejadian {kej_str} berlaku di {alamat_str}, {daerah_str}.{status_str}"
            else:
                full_event_sentence = f"{prefix}kejadian {kej_str} di {alamat_str}, {daerah_str}.{status_str}"
                 
            narrative_parts.append(full_event_sentence)
             
        h41_text += "".join(narrative_parts) + penutup_text
 
    apply_font(h41.add_run(h41_text), 11, bold=False)
 
    add_table_title(doc, "Jadual 4", "Senarai Kejadian Insiden Bencana, Kecemasan dan Krisis (BKK)")
    t4 = doc.add_table(rows=len(bkk_table_df) + 1, cols=len(bkk_table_df.columns))
    t4.style = 'Table Grid'
    t4.width = content_width 
     
    h4_col_count = len(bkk_table_df.columns)
    for i, col in enumerate(bkk_table_df.columns):
        cell = t4.rows[0].cells[i]
        if i < h4_col_count-1: 
            set_cell_background(cell, "BFDFFF")
        else: 
            set_cell_background(cell, "FFFF00") 
             
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_text = str(col).strip()
        if header_text.upper() == "INSIDEN/BENCANA": header_text = "Insiden/Bencana"
        apply_font(p.add_run(header_text.replace(" ", "\n")), 8, bold=True)
     
    for r_idx, row_data in enumerate(bkk_table_df.values):
        cells = t4.rows[r_idx+1].cells
        is_last_row = (r_idx == len(bkk_table_df)-1)
        for c_idx, val in enumerate(row_data):
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_val(val))
            apply_font(run, 8, bold=is_last_row or c_idx == 0)
            if is_last_row: 
                set_cell_background(cells[c_idx], "FFFF00")
            elif c_idx == 0: 
                set_cell_background(cells[c_idx], "D9E9FF")
            elif c_idx == bkk_table_df.shape[1]-1: 
                set_cell_background(cells[c_idx], "FFFFB3") 
 
    p4_intro = doc.add_paragraph()
    apply_font(p4_intro.add_run(" "), 11)
 
    p4_head = doc.add_paragraph()
    apply_font(p4_head.add_run("5.0 Lain-lain (Input secara manual)"), 11, bold=True)
     
    doc.add_paragraph() 
    sig_table = doc.add_table(rows=11, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    sig_table.columns[0].width = Cm(3.2)
    sig_table.columns[1].width = Cm(0.5)
    sig_table.columns[2].width = Cm(10)
 
    def fill_sig_row(row_idx, label):
        row = sig_table.rows[row_idx].cells
        p_label = row[0].paragraphs[0]
        apply_font(p_label.add_run(label), 11, bold=False)
        p_colon = row[1].paragraphs[0]
        apply_font(p_colon.add_run(":"), 11, bold=False)
 
    fill_sig_row(0, "Disediakan oleh")
    fill_sig_row(1, "Jawatan")
    fill_sig_row(2, "Tarikh")
    sig_table.rows[3].height = Pt(25)
 
    fill_sig_row(4, "Disemak")
    fill_sig_row(5, "Jawatan")
    fill_sig_row(6, "Tarikh")
    sig_table.rows[7].height = Pt(25)
 
    fill_sig_row(8, "Disahkan")
    fill_sig_row(9, "Jawatan")
    fill_sig_row(10, "Tarikh")
 
    tbl = sig_table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
        tbl.insert(0, tblPr)
    tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr.append(tblBorders)
 
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target
 
# --- STREAMLIT UI ---
st.set_page_config(page_title="BWKK Report Generator", layout="centered")
st.title("📋 BWKK Report Generator")
 
now_msia = get_msia_time()
today = now_msia.date()
yesterday = today - timedelta(days=1)
tarikh_guideline = get_malay_date(yesterday)
 
st.subheader("📁 Muat Naik Excel Notifikasi Harian")
st.info(f"**Guideline:** Sila muat turun file notifikasi pada **{tarikh_guideline}** dari sistem eNotifikasi dan muat naik di sini.")
f1 = st.file_uploader("Pilih fail Notifikasi Harian", type=["xlsx", "xls"], label_visibility="collapsed")
 
st.markdown("---")
 
st.subheader("📁 Muat Naik Excel Linelisting Wabak")
st.info("**Guideline:** Sila muat turun file google sheet **'LAPORAN WABAK NEGERI SELANGOR 3'** melalui email **cprc_sel@moh.gov.my** dan muat naik di sini.")
f2 = st.file_uploader("Pilih fail Linelisting Wabak", type=["xlsx", "xls"], label_visibility="collapsed")
 
if f1 and f2:
    if st.button("🚀 Jana Laporan Lengkap"):
        try:
            engine_type = "xlrd" if f1.name.endswith(".xls") else "openpyxl"
            df1 = pd.read_excel(f1, engine=engine_type)
            df1 = df1[df1['Notifikasi Status'] != 'Abai Notifikasi']
            df1 = df1[df1['Pejabat Kesihatan'].isin(TEMPLATE_PKDS)]
            matrix = pd.crosstab(df1['Diagnosis'], df1['Pejabat Kesihatan']).reindex(columns=TEMPLATE_PKDS, fill_value=0)
            matrix['Grand Total'] = matrix.sum(axis=1)
            matrix['Average Harian'] = [AVG_HARIAN_FIGURES.get(format_penyakit_name(idx), 0) for idx in matrix.index]
            matrix = matrix.sort_values(by='Grand Total', ascending=False)
            col_totals = matrix[TEMPLATE_PKDS + ['Grand Total']].sum(axis=0)
 
            engine_type2 = "xlrd" if f2.name.endswith(".xls") else "openpyxl"
            df2 = pd.read_excel(f2, sheet_name="SELANGOR 2", engine=engine_type2)
            df2['Tarikh Isytihar Wabak'] = pd.to_datetime(df2['Tarikh Isytihar Wabak']).dt.date
            df2['Tarikh Sebenar Tamat Wabak'] = pd.to_datetime(df2['Tarikh Sebenar Tamat Wabak '], errors='coerce').dt.date
            df2['Tarikh Wabak Dijangka Tamat'] = pd.to_datetime(df2['Tarikh Wabak Dijangka Tamat'], errors='coerce').dt.date
 
            addr_col = 'Tempat Berlaku Wabak\n(Alamat diisi lengkap dengan :- No rumah, nama jalan, nama tempat, daerah dan Negeri)'
            cat_col = 'Kategori Tempat\n(Kategori premis berdasarkan tempat berlaku wabak)'
            df_yesterday = df2[df2['Tarikh Isytihar Wabak'] == yesterday].copy()
            df_yesterday_list = df_yesterday[['PENYAKIT', 'DAERAH (HURUF BESAR)', addr_col, cat_col, 'Bilangan Kes', 'Bilangan Terdedah']].values.tolist()
 
            df2_filt = df2[df2['Tarikh Isytihar Wabak'] >= date(2026, 1, 4)].copy()
            def group_inf(n): return "ILI/ Influenza" if any(x in str(n).upper() for x in ["INFLUENZA", "ILI"]) else n
            df2_filt['PENYAKIT'] = df2_filt['PENYAKIT'].apply(group_inf)
             
            wb_sum = []
            for d in df2_filt['PENYAKIT'].unique():
                if pd.isna(d): continue
                disease_df = df2_filt[df2_filt['PENYAKIT'] == d]
                h = len(disease_df[disease_df['Tarikh Isytihar Wabak'] == yesterday])
                k = len(disease_df)
                def check_active(row):
                    tamat = row['Tarikh Sebenar Tamat Wabak'] if pd.notna(row['Tarikh Sebenar Tamat Wabak']) else row['Tarikh Wabak Dijangka Tamat']
                    return True if (pd.isna(tamat) or tamat >= today) else False
                active_count = disease_df.apply(check_active, axis=1).sum()
                wb_sum.append({'PENYAKIT': d, 'HARIAN': h, 'AKTIF': active_count, 'KUMULATIF': k})
            wabak_df = pd.DataFrame(wb_sum).set_index('PENYAKIT').sort_values(by='KUMULATIF', ascending=False)
 
            raw_gs = pd.read_csv(GSHEET_URL, header=None)
            mask_v = raw_gs.apply(lambda r: r.astype(str).str.contains('Petaling').any(), axis=1)
            v_data = raw_gs.iloc[mask_v.idxmax() : mask_v.idxmax() + 11, 13:20]
            v_data = v_data.dropna(how='all')
            v_data = v_data[~v_data.iloc[:, 0].astype(str).str.lower().str.contains('nan')]
 
            # --- PEMPROSESAN DATA GOOGLE SHEET BKK ---
            df_bkk_raw_data = pd.read_csv(URL_BKK_LINELISTING, header=None)
            clean_date_series = df_bkk_raw_data.iloc[:, 2].astype(str).str.strip()
            df_bkk_raw_data['datetime_lapor'] = pd.to_datetime(clean_date_series, dayfirst=True, errors='coerce').dt.date
             
            insiden_semalam = df_bkk_raw_data[df_bkk_raw_data['datetime_lapor'] == yesterday]
             
            bkk_details = [{
                'kejadian': r[5], 
                'alamat': r[8], 
                'daerah': r[4],
                'bil_kes': r[9],      
                'bil_kematian': r[10]  
            } for _, r in insiden_semalam.iterrows()]
             
            df_bkk_jadual_full = pd.read_csv(URL_BKK_JADUAL, header=None)
             
            bkk_raw = df_bkk_jadual_full.iloc[1:, 33:46].dropna(how='all').reset_index(drop=True)
            bkk_raw.columns = bkk_raw.iloc[0]
             
            # --- PEMPROSESAN INSENSITIF KES NAMA KOLUM ---
            new_cols = []
            for col in bkk_raw.columns:
                col_str = str(col).strip()
                if re.search(r'(moving|median|4 tahun)', col_str, re.IGNORECASE):
                    new_cols.append('Purata Bergerak 4 Tahun (2022,2023,2024,2025)')
                else:
                    new_cols.append(col_str)
            bkk_raw.columns = new_cols
 
            bkk_table_final = bkk_raw[1:].reset_index(drop=True).rename(columns={
                'GOMBAK':'GBK','HULU LANGAT':'HL','HULU SELANGOR':'HS','KLANG':'KLG','KUALA LANGAT':'KL','KUALA SELANGOR':'KS','PETALING':'PTG','SABAK BERNAM':'SB','SEPANG':'SPG'
            })
 
            # --- BAGIAN NOTA DIAGNOSTIK DI SINI TELAH DIBUANG ---
 
            doc_out = generate_docx(matrix, col_totals, wabak_df, v_data, bkk_table_final, (len(bkk_details)==0), bkk_details, df_yesterday_list)
             
            file_date = today.strftime("%d.%m.%y")
            file_name_custom = f"LAPORAN CPRC ({file_date}).docx"
 
            st.success(f"✅ Laporan berjaya dijana untuk tarikh {file_date}!")
            st.download_button(
                label="⬇️ Muat Turun Laporan", 
                data=doc_out, 
                file_name=file_name_custom,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
 
        except Exception as e:
            st.error(f"Ralat: {e}")
